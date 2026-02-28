"""Конвертер Food_Data_TZ.md → Food_Data_TZ.docx"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

MD_PATH = Path(__file__).parent / "Food_Data_TZ.md"
DOCX_PATH = Path(__file__).parent / "Food_Data_TZ.docx"


def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)


def add_formatted_run(paragraph, text: str, bold=False, italic=False, font_name='Calibri', font_size=11, color=None, is_code=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if is_code:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return run


def parse_inline(paragraph, text: str, font_size=11, base_bold=False):
    """Парсит inline markdown: **bold**, `code`, обычный текст."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            add_formatted_run(paragraph, part[2:-2], bold=True, font_size=font_size)
        elif part.startswith('`') and part.endswith('`'):
            add_formatted_run(paragraph, part[1:-1], is_code=True)
        else:
            add_formatted_run(paragraph, part, bold=base_bold, font_size=font_size)


def add_table(doc, rows_data: list[list[str]]):
    if not rows_data or len(rows_data) < 2:
        return
    
    n_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    for ci in range(n_cols):
        for ri, row in enumerate(rows_data):
            cell = table.cell(ri, ci)
            cell_text = row[ci] if ci < len(row) else ''
            
            for p in cell.paragraphs:
                p.clear()
            
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            
            if ri == 0:
                set_cell_shading(cell, '2B579A')
                add_formatted_run(para, cell_text.strip(), bold=True, font_size=9, color=(255, 255, 255))
            else:
                parse_inline(para, cell_text.strip(), font_size=9)


def convert():
    md = MD_PATH.read_text(encoding='utf-8')
    lines = md.split('\n')
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1A, 0x47, 0x80)
    
    doc.styles['Heading 1'].font.size = Pt(18)
    doc.styles['Heading 2'].font.size = Pt(15)
    doc.styles['Heading 3'].font.size = Pt(13)
    doc.styles['Heading 4'].font.size = Pt(11)
    
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    
    i = 0
    table_buffer = []
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                add_formatted_run(para, code_text, is_code=True)
                i += 1
                continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        if line.startswith('|') and '|' in line[1:]:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            
            if i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1]):
                if table_buffer:
                    add_table(doc, table_buffer)
                    table_buffer = []
                table_buffer.append(cells)
                i += 2
                continue
            elif table_buffer:
                table_buffer.append(cells)
                i += 1
                if i >= len(lines) or not lines[i].startswith('|'):
                    add_table(doc, table_buffer)
                    table_buffer = []
                continue
        else:
            if table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
        
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue
        
        if line.strip() == '---':
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        
        if line.strip() == '':
            i += 1
            continue
        
        if re.match(r'^- \[ \] ', line):
            para = doc.add_paragraph(style='List Bullet')
            text = line[6:]
            add_formatted_run(para, '☐ ', font_size=11)
            parse_inline(para, text)
            i += 1
            continue
        
        if line.startswith('- '):
            para = doc.add_paragraph(style='List Bullet')
            parse_inline(para, line[2:])
            i += 1
            continue
        
        if re.match(r'^\d+\.\s', line):
            para = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', line)
            parse_inline(para, text)
            i += 1
            continue
        
        para = doc.add_paragraph()
        parse_inline(para, line)
        i += 1
    
    if table_buffer:
        add_table(doc, table_buffer)
    
    doc.save(str(DOCX_PATH))
    print(f"Saved: {DOCX_PATH}")


if __name__ == '__main__':
    convert()
